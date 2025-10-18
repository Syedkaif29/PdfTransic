from fastapi import APIRouter, HTTPException, UploadFile, File, Form
from fastapi.responses import StreamingResponse
from typing import List, Dict, Any, AsyncGenerator
import json
import io
import logging
from dataclasses import asdict
from font_management import get_language_mapper
from translation_core import extract_text_from_pdf, create_pdf_from_text, chunk_text, remove_duplicates_from_text, ip, tokenizer, model, DEVICE, PdfLayoutData
import torch
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch


router = APIRouter()
logger = logging.getLogger(__name__)
logger.setLevel(logging.DEBUG)



@router.post("/translate-pdf-enhanced")
async def translate_pdf_enhanced(
    file: UploadFile = File(...),
    target_language: str = Form(...)
):
    """Enhanced PDF translation with duplicate removal and layout data"""
    global tokenizer, model, ip, DEVICE
    if not all([tokenizer, model, ip]):
        raise HTTPException(status_code=503, detail="Models are still loading. Please try again in a moment.")
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Only PDF files are supported")
    
    try:
        pdf_content = await file.read()
        extracted_text, extraction_method, layout_data = extract_text_from_pdf(pdf_content, max_pages=2)
        
        if not extracted_text.strip():
            raise HTTPException(status_code=422, detail="Could not extract text from PDF")
        
        # Remove duplicates for cleaner translation
        cleaned_text = remove_duplicates_from_text(extracted_text)
        duplicates_removed = len(extracted_text) != len(cleaned_text)
        
        # Split into chunks for processing
        text_chunks = chunk_text(cleaned_text, max_chunk_size=500)
        
        # Translate chunks in batches
        all_translations = []
        batch_size = 19
        
        for i in range(0, len(text_chunks), batch_size):
            batch = text_chunks[i:i+batch_size]
            try:
                pre = ip.preprocess_batch(batch, src_lang="eng_Latn", tgt_lang=target_language)
                inputs = tokenizer(
                    pre,
                    truncation=True,
                    padding="longest",
                    return_tensors="pt",
                    return_attention_mask=True,
                    max_length=512,
                ).to(DEVICE)
                
                with torch.no_grad():
                    generated_tokens = model.generate(
                        input_ids=inputs["input_ids"],
                        attention_mask=inputs["attention_mask"],
                        use_cache=False,
                        min_length=0,
                        max_length=256,
                        num_beams=3,
                        num_return_sequences=1,
                        do_sample=False,
                    )
                
                decoded = tokenizer.batch_decode(
                    generated_tokens.detach().cpu().tolist(),
                    skip_special_tokens=True,
                    clean_up_tokenization_spaces=True,
                )
                
                batch_trans = ip.postprocess_batch(decoded, lang=target_language)
                all_translations.extend(batch_trans)
                
            except Exception as batch_error:
                logger.error(f"Error translating batch: {batch_error}")
                all_translations.extend(["[Translation failed]"] * len(batch))
        
        # Join translations
        translated_text = ' '.join(all_translations)
        
        return {
            "success": True,
            "filename": file.filename,
            "pages_processed": 2,
            "extracted_text": extracted_text,
            "translated_text": translated_text,
            "target_language": target_language,
            "source_language": "eng_Latn",
            "extraction_method": extraction_method,
            "duplicates_removed": duplicates_removed,
            "text_length": len(extracted_text),
            "chunks_processed": len(text_chunks),
            "memory_management": "chunking",
            "layout_data_available": bool(layout_data),
            "original_text_chunks": text_chunks,
            "translated_text_chunks": all_translations,
            "layout_data": [asdict(p) for p in layout_data] if layout_data else []
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error in enhanced PDF translation: {e}")
        raise HTTPException(status_code=500, detail=f"Enhanced PDF translation failed: {e}")

@router.post("/translate-and-download-pdf")
async def translate_and_download_pdf(
    file: UploadFile = File(...),
    target_language: str = Form(...)
):
    global tokenizer, model, ip, DEVICE
    if not all([tokenizer, model, ip]):
        raise HTTPException(status_code=503, detail="Models are still loading. Please try again in a moment.")
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Only PDF files are supported")
    try:
        pdf_content = await file.read()
        extracted_text, extraction_method, layout_data = extract_text_from_pdf(pdf_content, max_pages=2)
        if not layout_data or not any(p.text_elements for p in layout_data):
            raise HTTPException(status_code=422, detail="Could not extract layout from PDF. Only text-based PDFs are supported for layout-preserving translation.")
        all_elements = []
        for page in layout_data:
            all_elements.extend(page.text_elements)
        seen = set()
        element_texts = []
        element_indices = []
        for idx, el in enumerate(all_elements):
            t = el.text.strip()
            if t and t not in seen:
                element_texts.append(t)
                element_indices.append(idx)
                seen.add(t)
            else:
                element_indices.append(None)
        batch_size = 19
        translated_texts = []
        for i in range(0, len(element_texts), batch_size):
            batch = element_texts[i:i+batch_size]
            try:
                pre = ip.preprocess_batch(batch, src_lang="eng_Latn", tgt_lang=target_language)
                inputs = tokenizer(
                    pre,
                    truncation=True,
                    padding="longest",
                    return_tensors="pt",
                    return_attention_mask=True,
                    max_length=512,
                ).to(DEVICE)
                with torch.no_grad():
                    generated_tokens = model.generate(
                        input_ids=inputs["input_ids"],
                        attention_mask=inputs["attention_mask"],
                        use_cache=False,
                        min_length=0,
                        max_length=256,
                        num_beams=3,
                        num_return_sequences=1,
                        do_sample=False,
                    )
                decoded = tokenizer.batch_decode(
                    generated_tokens.detach().cpu().tolist(),
                    skip_special_tokens=True,
                    clean_up_tokenization_spaces=True,
                )
                batch_trans = ip.postprocess_batch(decoded, lang=target_language)
                translated_texts.extend(batch_trans)
            except Exception as batch_error:
                logger.error(f"Error translating element batch: {batch_error}")
                translated_texts.extend(["[Translation failed]"] * len(batch))
        translated_elements = []
        trans_idx = 0
        for idx in element_indices:
            if idx is not None:
                translated_elements.append(translated_texts[trans_idx])
                trans_idx += 1
            else:
                translated_elements.append("")
        pdf_buffer = create_pdf_from_text(
            original_text_chunks=[el.text for el in all_elements],
            translated_text_chunks=translated_elements,
            filename=file.filename,
            target_language=target_language,
            language_code=target_language,
            layout_data=layout_data
        )
        pdf_buffer.seek(0)
        headers = {
            "Content-Disposition": f"attachment; filename=translated_{file.filename}"
        }
        return StreamingResponse(iter([pdf_buffer.getvalue()]), media_type="application/pdf", headers=headers)
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error in direct PDF translation: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to generate translated PDF: {e}")

@router.post("/translate-pdf-live-preview")
async def translate_pdf_live_preview(
    file: UploadFile = File(...),
    target_language: str = Form(...)
):
    global tokenizer, model, ip, DEVICE
    if not all([tokenizer, model, ip]):
        raise HTTPException(status_code=503, detail="Models are still loading. Please try again in a moment.")
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Only PDF files are supported")
    pdf_content = await file.read()
    extracted_text, extraction_method, layout_data = extract_text_from_pdf(pdf_content, max_pages=2)
    if not layout_data or not any(p.text_elements for p in layout_data):
        raise HTTPException(status_code=422, detail="Could not extract layout from PDF. Only text-based PDFs are supported for layout-preserving translation.")
    all_elements = []
    for page in layout_data:
        all_elements.extend(page.text_elements)
    seen = set()
    element_texts = []
    element_indices = []
    for idx, el in enumerate(all_elements):
        t = el.text.strip()
        if t and t not in seen:
            element_texts.append(t)
            element_indices.append(idx)
            seen.add(t)
        else:
            element_indices.append(None)
    async def event_generator() -> AsyncGenerator[str, None]:
        # First, send layout information
        layout_info_event = {
            "type": "layout_info",
            "total_elements": len(element_texts),
            "layout_data": [
                {
                    "index": idx,
                    "text": el.text,
                    "bbox": list(el.bbox),
                    "font_name": el.font_name,
                    "font_size": el.font_size,
                    "color": list(el.color),
                    "is_bold": el.is_bold,
                    "is_italic": el.is_italic,
                    "paragraph_index": idx // 10,  # Simple paragraph grouping
                    "word_index": idx,
                    "is_title": el.font_size > 16,
                    "is_heading": el.font_size > 14 and el.font_size <= 16
                }
                for idx, el in enumerate(all_elements) if element_indices[idx] is not None
            ]
        }
        yield f"data: {json.dumps(layout_info_event)}\n\n"
        
        batch_size = 19
        trans_idx = 0
        total_elements = len(element_texts)
        
        for i in range(0, len(element_texts), batch_size):
            batch = element_texts[i:i+batch_size]
            try:
                pre = ip.preprocess_batch(batch, src_lang="eng_Latn", tgt_lang=target_language)
                inputs = tokenizer(
                    pre,
                    truncation=True,
                    padding="longest",
                    return_tensors="pt",
                    return_attention_mask=True,
                    max_length=512,
                ).to(DEVICE)
                with torch.no_grad():
                    generated_tokens = model.generate(
                        input_ids=inputs["input_ids"],
                        attention_mask=inputs["attention_mask"],
                        use_cache=False,
                        min_length=0,
                        max_length=256,
                        num_beams=3,
                        num_return_sequences=1,
                        do_sample=False,
                    )
                decoded = tokenizer.batch_decode(
                    generated_tokens.detach().cpu().tolist(),
                    skip_special_tokens=True,
                    clean_up_tokenization_spaces=True,
                )
                batch_trans = ip.postprocess_batch(decoded, lang=target_language)
            except Exception as batch_error:
                logger.error(f"Error translating element batch: {batch_error}")
                batch_trans = ["[Translation failed]"] * len(batch)
            
            for j, orig in enumerate(batch):
                # Find the corresponding layout element
                layout_element = None
                actual_element_idx = None
                for idx, el in enumerate(all_elements):
                    if element_indices[idx] == trans_idx and el.text.strip() == orig:
                        layout_element = {
                            "index": idx,
                            "text": el.text,
                            "bbox": list(el.bbox),
                            "font_name": el.font_name,
                            "font_size": el.font_size,
                            "color": list(el.color),
                            "is_bold": el.is_bold,
                            "is_italic": el.is_italic,
                            "paragraph_index": idx // 10,
                            "word_index": idx,
                            "is_title": el.font_size > 16,
                            "is_heading": el.font_size > 14 and el.font_size <= 16
                        }
                        actual_element_idx = idx
                        break
                
                event = {
                    "type": "translation_update",
                    "element_index": trans_idx,
                    "original_text": orig,
                    "translated_text": batch_trans[j],
                    "layout": layout_element,
                    "progress": {
                        "current": trans_idx + 1,
                        "total": total_elements,
                        "percentage": ((trans_idx + 1) / total_elements) * 100,
                        "batch_progress": f"Batch {i//batch_size + 1}/{(len(element_texts) + batch_size - 1)//batch_size}"
                    }
                }
                yield f"data: {json.dumps(event)}\n\n"
                trans_idx += 1
        
        # Send completion event
        completion_event = {
            "type": "translation_complete",
            "total_translated": trans_idx,
            "message": "Translation completed successfully"
        }
        yield f"data: {json.dumps(completion_event)}\n\n"
    return StreamingResponse(event_generator(), media_type="text/event-stream")

@router.post("/translate-batch")
async def translate_batch(
    files: List[UploadFile] = File(...),
    target_language: str = Form(...)
):
    """Batch translate multiple PDF files"""
    global tokenizer, model, ip, DEVICE
    if not all([tokenizer, model, ip]):
        raise HTTPException(status_code=503, detail="Models are still loading. Please try again in a moment.")
    
    if len(files) > 5:  # Limit batch size
        raise HTTPException(status_code=400, detail="Maximum 5 files allowed per batch")
    
    results = []
    
    for i, file in enumerate(files):
        if not file.filename.lower().endswith('.pdf'):
            results.append({
                "filename": file.filename,
                "success": False,
                "error": "Only PDF files are supported"
            })
            continue
            
        try:
            pdf_content = await file.read()
            extracted_text, extraction_method, layout_data = extract_text_from_pdf(pdf_content, max_pages=2)
            
            if not extracted_text.strip():
                results.append({
                    "filename": file.filename,
                    "success": False,
                    "error": "Could not extract text from PDF"
                })
                continue
            
            # Process translation (simplified for batch)
            cleaned_text = remove_duplicates_from_text(extracted_text)
            text_chunks = chunk_text(cleaned_text, max_chunk_size=300)  # Smaller chunks for batch
            
            all_translations = []
            batch_size = 19
            
            for j in range(0, len(text_chunks), batch_size):
                batch = text_chunks[j:j+batch_size]
                try:
                    pre = ip.preprocess_batch(batch, src_lang="eng_Latn", tgt_lang=target_language)
                    inputs = tokenizer(
                        pre,
                        truncation=True,
                        padding="longest",
                        return_tensors="pt",
                        return_attention_mask=True,
                        max_length=256,  # Reduced for batch processing
                    ).to(DEVICE)
                    
                    with torch.no_grad():
                        generated_tokens = model.generate(
                            input_ids=inputs["input_ids"],
                            attention_mask=inputs["attention_mask"],
                            use_cache=False,
                            min_length=0,
                            max_length=128,  # Reduced for batch processing
                            num_beams=2,  # Reduced for speed
                            num_return_sequences=1,
                            do_sample=False,
                        )
                    
                    decoded = tokenizer.batch_decode(
                        generated_tokens.detach().cpu().tolist(),
                        skip_special_tokens=True,
                        clean_up_tokenization_spaces=True,
                    )
                    
                    batch_trans = ip.postprocess_batch(decoded, lang=target_language)
                    all_translations.extend(batch_trans)
                    
                except Exception as batch_error:
                    logger.error(f"Error in batch translation: {batch_error}")
                    all_translations.extend(["[Translation failed]"] * len(batch))
            
            translated_text = ' '.join(all_translations)
            
            results.append({
                "filename": file.filename,
                "success": True,
                "extracted_text": extracted_text[:500] + "..." if len(extracted_text) > 500 else extracted_text,
                "translated_text": translated_text,
                "target_language": target_language,
                "extraction_method": extraction_method,
                "chunks_processed": len(text_chunks),
                "file_index": i + 1,
                "total_files": len(files)
            })
            
        except Exception as e:
            logger.error(f"Error processing file {file.filename}: {e}")
            results.append({
                "filename": file.filename,
                "success": False,
                "error": str(e)
            })
    
    return {
        "batch_results": results,
        "total_files": len(files),
        "successful_translations": len([r for r in results if r.get("success", False)]),
        "failed_translations": len([r for r in results if not r.get("success", False)]),
        "target_language": target_language
    }

@router.post("/translate-and-download-docx")
async def translate_and_download_docx(
    file: UploadFile = File(...),
    target_language: str = Form(...)
):
    global tokenizer, model, ip, DEVICE
    if not all([tokenizer, model, ip]):
        raise HTTPException(status_code=503, detail="Models are still loading. Please try again in a moment.")
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Only PDF files are supported")
    try:
        pdf_content = await file.read()
        extracted_text, extraction_method, layout_data = extract_text_from_pdf(pdf_content, max_pages=2)
        if not layout_data or not any(p.text_elements for p in layout_data):
            raise HTTPException(status_code=422, detail="Could not extract layout from PDF. Only text-based PDFs are supported for layout-preserving translation.")
        all_elements = []
        for page in layout_data:
            all_elements.extend(page.text_elements)
        seen = set()
        element_texts = []
        element_indices = []
        for idx, el in enumerate(all_elements):
            t = el.text.strip()
            if t and t not in seen:
                element_texts.append(t)
                element_indices.append(idx)
                seen.add(t)
            else:
                element_indices.append(None)
        batch_size = 19
        translated_texts = []
        for i in range(0, len(element_texts), batch_size):
            batch = element_texts[i:i+batch_size]
            try:
                pre = ip.preprocess_batch(batch, src_lang="eng_Latn", tgt_lang=target_language)
                inputs = tokenizer(
                    pre,
                    truncation=True,
                    padding="longest",
                    return_tensors="pt",
                    return_attention_mask=True,
                    max_length=512,
                ).to(DEVICE)
                with torch.no_grad():
                    generated_tokens = model.generate(
                        input_ids=inputs["input_ids"],
                        attention_mask=inputs["attention_mask"],
                        use_cache=False,
                        min_length=0,
                        max_length=256,
                        num_beams=3,
                        num_return_sequences=1,
                        do_sample=False,
                    )
                decoded = tokenizer.batch_decode(
                    generated_tokens.detach().cpu().tolist(),
                    skip_special_tokens=True,
                    clean_up_tokenization_spaces=True,
                )
                batch_trans = ip.postprocess_batch(decoded, lang=target_language)
                translated_texts.extend(batch_trans)
            except Exception as batch_error:
                logger.error(f"Error translating element batch: {batch_error}")
                translated_texts.extend(["[Translation failed]"] * len(batch))
        translated_elements = []
        trans_idx = 0
        for idx in element_indices:
            if idx is not None:
                translated_elements.append(translated_texts[trans_idx])
                trans_idx += 1
            else:
                translated_elements.append("")
        # Build the .docx file
        doc = Document()
        doc.add_heading(f"Translated PDF: {file.filename}", 0)
        last_font_size = None
        for i, el in enumerate(all_elements):
            translated_text = translated_elements[i] if i < len(translated_elements) else ""
            if not translated_text.strip():
                continue
            p = doc.add_paragraph()
            run = p.add_run(translated_text)
            # Font size
            font_size = int(el.font_size) if hasattr(el, 'font_size') and el.font_size else 12
            run.font.size = Pt(font_size)
            # Bold/Italic
            run.bold = getattr(el, 'is_bold', False)
            run.italic = getattr(el, 'is_italic', False)
            # Color (convert 0-1 RGB to 0-255)
            if hasattr(el, 'color') and el.color:
                r, g, b = [int(255 * c) for c in el.color]
                run.font.color.rgb = RGBColor(r, g, b)
            # Alignment (left by default)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        headers = {
            "Content-Disposition": f"attachment; filename=translated_{file.filename.replace('.pdf', '.docx')}"
        }
        return StreamingResponse(iter([buffer.getvalue()]), media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers=headers)
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error in direct DOCX translation: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to generate translated DOCX: {e}")

@router.post("/download-translated-pdf")
async def download_translated_pdf(
    original_text_chunks_json: str = Form(...),
    translated_text_chunks_json: str = Form(...),
    layout_data_json: str = Form(...),
    filename: str = Form(...),
    target_language: str = Form(...)
):
    """Download translated document as PDF format from Extract & Translate"""
    logger.info(f"PDF download request received")
    logger.info(f"Filename: {filename}")
    logger.info(f"Target language: {target_language}")
    
    try:
        # Parse the JSON data
        original_text_chunks = json.loads(original_text_chunks_json)
        translated_text_chunks = json.loads(translated_text_chunks_json)
        layout_data_dicts = json.loads(layout_data_json)
        
        logger.info(f"Parsed data: {len(translated_text_chunks)} translated chunks")
        
        # Validate data
        if not translated_text_chunks:
            raise HTTPException(status_code=422, detail="No translated text chunks provided")
        
        # Convert layout_data dictionaries back to PdfLayoutData objects
        layout_data = None
        if layout_data_dicts:
            try:
                layout_data = [PdfLayoutData(**item) for item in layout_data_dicts]
                logger.info(f"Converted {len(layout_data)} layout data objects")
            except Exception as e:
                logger.warning(f"Failed to convert layout data: {e}. Proceeding without layout.")
                layout_data = None
        
        # Create PDF using the existing create_pdf_from_text function
        pdf_buffer = create_pdf_from_text(
            original_text_chunks=original_text_chunks,
            translated_text_chunks=translated_text_chunks,
            filename=filename,
            target_language=target_language,
            language_code=target_language,
            layout_data=layout_data
        )
        
        pdf_buffer.seek(0)
        
        # Generate filename
        pdf_filename = filename.replace('.PDF', '.pdf')
        if not pdf_filename.endswith('.pdf'):
            pdf_filename += '.pdf'
        
        logger.info(f"Successfully created PDF: {len(pdf_buffer.getvalue())} bytes")
        
        return StreamingResponse(
            iter([pdf_buffer.getvalue()]),
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename=translated_{pdf_filename}"}
        )
        
    except Exception as e:
        logger.error(f"Error creating PDF: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error creating PDF document: {str(e)}")

@router.post("/download-translated-pdf-simple")
async def download_translated_pdf_simple(
    file: UploadFile = File(...),
    filename: str = Form(...),
    target_language: str = Form(...),
    original_text_chunks_json: str = Form(...),
    translated_text_chunks_json: str = Form(...),
    layout_data_json: str = Form(...),
    language_code: str = Form(...)
):
    """Download translated document as PDF format from live preview - simplified version"""
    logger.info(f"Simple PDF download request received")
    logger.info(f"Filename: {filename}")
    logger.info(f"Target language: {target_language}")
    
    try:
        # Parse the JSON data (same as Word endpoint)
        original_text_chunks = json.loads(original_text_chunks_json)
        translated_text_chunks = json.loads(translated_text_chunks_json)
        layout_data = json.loads(layout_data_json)
        
        logger.info(f"Parsed data: {len(translated_text_chunks)} translated chunks")
        
        # Validate data
        if not translated_text_chunks:
            raise HTTPException(status_code=422, detail="No translated text chunks provided")
        
        # Create PDF using ReportLab (simple approach)
        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=A4)
        width, height = A4
        margin = 1 * inch
        y_position = height - margin
        line_height = 14
        
        # Title
        c.setFont("Helvetica-Bold", 20)
        title_text = "Translated Document"
        title_width = c.stringWidth(title_text, "Helvetica-Bold", 20)
        c.drawString((width - title_width) / 2, y_position, title_text)
        y_position -= 2 * line_height
        
        # Language info
        c.setFont("Helvetica", 12)
        lang_text = f"Translated to: {target_language}"
        lang_width = c.stringWidth(lang_text, "Helvetica", 12)
        c.drawString((width - lang_width) / 2, y_position, lang_text)
        y_position -= line_height
        
        # Separator
        separator = "─" * 50
        sep_width = c.stringWidth(separator, "Helvetica", 12)
        c.drawString((width - sep_width) / 2, y_position, separator)
        y_position -= 2 * line_height
        
        # Content - just add all translated text
        c.setFont("Helvetica", 12)
        all_translated = ' '.join(translated_text_chunks)
        
        # Simple text wrapping
        words = all_translated.split()
        current_line = ""
        
        for word in words:
            test_line = current_line + " " + word if current_line else word
            if c.stringWidth(test_line, "Helvetica", 12) < (width - 2 * margin):
                current_line = test_line
            else:
                if current_line:
                    if y_position < margin:
                        c.showPage()
                        y_position = height - margin
                    c.drawString(margin, y_position, current_line)
                    y_position -= line_height
                current_line = word
        
        # Draw remaining text
        if current_line:
            if y_position < margin:
                c.showPage()
                y_position = height - margin
            c.drawString(margin, y_position, current_line)
        
        c.save()
        buffer.seek(0)
        
        # Generate filename
        pdf_filename = filename.replace('.PDF', '.pdf')
        if not pdf_filename.endswith('.pdf'):
            pdf_filename += '.pdf'
        
        logger.info(f"Successfully created simple PDF: {len(buffer.getvalue())} bytes")
        
        return StreamingResponse(
            iter([buffer.getvalue()]),
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename=translated_{pdf_filename}"}
        )
        
    except Exception as e:
        logger.error(f"Error creating simple PDF: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error creating PDF document: {str(e)}")

@router.post("/analyze-translation-quality")
async def analyze_translation_quality(
    original_text: str = Form(...),
    translated_text: str = Form(...),
    target_language: str = Form(...)
):
    """Analyze translation quality and provide confidence metrics"""
    try:
        # Simple quality metrics (can be enhanced with more sophisticated algorithms)
        original_words = len(original_text.split())
        translated_words = len(translated_text.split())
        
        # Length ratio analysis
        length_ratio = translated_words / original_words if original_words > 0 else 0
        length_score = 100 - abs(length_ratio - 1) * 50  # Penalize extreme length differences
        
        # Character diversity analysis
        original_chars = set(original_text.lower())
        translated_chars = set(translated_text.lower())
        diversity_score = len(translated_chars) / len(original_chars) * 100 if len(original_chars) > 0 else 0
        
        # Simple completeness check
        completeness_score = min(100, (len(translated_text.strip()) / len(original_text.strip())) * 100) if len(original_text.strip()) > 0 else 0
        
        # Overall confidence (weighted average)
        confidence = (length_score * 0.3 + diversity_score * 0.3 + completeness_score * 0.4)
        confidence = max(60, min(95, confidence))  # Clamp between 60-95%
        
        return {
            "confidence": round(confidence, 1),
            "metrics": {
                "length_ratio": round(length_ratio, 2),
                "length_score": round(length_score, 1),
                "diversity_score": round(diversity_score, 1),
                "completeness_score": round(completeness_score, 1),
                "original_words": original_words,
                "translated_words": translated_words
            },
            "recommendations": [
                "Translation appears complete" if completeness_score > 80 else "Translation may be incomplete",
                "Good length balance" if 0.7 <= length_ratio <= 1.5 else "Unusual length difference detected",
                "High confidence translation" if confidence > 85 else "Consider manual review"
            ]
        }
        
    except Exception as e:
        logger.error(f"Error analyzing translation quality: {e}")
        raise HTTPException(status_code=500, detail=f"Quality analysis failed: {e}")

@router.post("/download-translated-word")
async def download_translated_word(
    file: UploadFile = File(...),
    filename: str = Form(...),
    target_language: str = Form(...),
    original_text_chunks_json: str = Form(...),
    translated_text_chunks_json: str = Form(...),
    layout_data_json: str = Form(...),
    language_code: str = Form(...)
):
    """Download translated document as Word format from live preview"""
    try:
        # Parse the JSON data
        original_text_chunks = json.loads(original_text_chunks_json)
        translated_text_chunks = json.loads(translated_text_chunks_json)
        layout_data = json.loads(layout_data_json)
        
        # Create Word document
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Pt(72)  # 1 inch
            section.bottom_margin = Pt(72)
            section.left_margin = Pt(72)
            section.right_margin = Pt(72)
        
        # Add title
        title = doc.add_heading('Translation Document', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add metadata
        metadata = doc.add_paragraph()
        metadata.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = metadata.add_run(f'Original File: {filename} | Target Language: {target_language}')
        run.font.size = Pt(10)
        
        # Add separator
        doc.add_paragraph('─' * 70)
        
        # ===== ORIGINAL TEXT SECTION =====
        original_heading = doc.add_heading('Original Text:', level=1)
        original_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        # Join all original chunks
        full_original = ' '.join(original_text_chunks)
        if full_original.strip():
            original_para = doc.add_paragraph(full_original)
            original_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            # Set font for original text
            for run in original_para.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
        
        # Add spacing between sections
        doc.add_paragraph()
        
        # Add separator
        doc.add_paragraph('─' * 70)
        
        # ===== TRANSLATED TEXT SECTION =====
        translated_heading = doc.add_heading('Translated Text:', level=1)
        translated_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        # Join all translated chunks
        full_translated = ' '.join(translated_text_chunks)
        if full_translated.strip():
            translated_para = doc.add_paragraph(full_translated)
            translated_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            # Set font for translated text (will use system font for Hindi/Urdu/etc.)
            for run in translated_para.runs:
                run.font.size = Pt(11)
        
        # Save to BytesIO
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        # Generate filename
        word_filename = filename.replace('.pdf', '.docx').replace('.PDF', '.docx')
        if not word_filename.endswith('.docx'):
            word_filename += '.docx'
        
        return StreamingResponse(
            io.BytesIO(doc_buffer.getvalue()),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=translated_{word_filename}"}
        )
        
    except Exception as e:
        logger.error(f"Error creating Word document: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error creating Word document: {str(e)}")
